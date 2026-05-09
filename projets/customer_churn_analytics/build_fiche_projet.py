"""
build_fiche_projet.py — Genere la fiche projet Customer Churn Analytics au format .docx.

Lance :
    pip install python-docx
    python build_fiche_projet.py

Sortie : fiche_projet_customer_churn_analytics.docx
"""

from __future__ import annotations

from pathlib import Path
import json

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ============================================================
# Palette IvoirCom — Orange chaleureux
# ============================================================
COLOR_TITLE_BG     = "1E3A5F"  # bandeau titre sombre (navy DataProjectLab)
COLOR_PRIMARY      = "E94E1B"  # orange chaleureux (sidebar / accents)
COLOR_PRIMARY_DARK = "A8350E"  # orange foncé
COLOR_SECONDARY    = "1D9E75"  # vert pastille compétences
COLOR_TEXT         = "2C2C2A"
COLOR_LABEL        = "888780"
COLOR_PAGE_BG      = "F9F9F8"
COLOR_BLOCK_BG     = "FFF4ED"  # orange très pâle (cards)
COLOR_BORDER       = "E5DDD5"

# ============================================================
# Helpers
# ============================================================

def set_cell_bg(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_run_color(run, color_hex: str) -> None:
    run.font.color.rgb = RGBColor.from_string(color_hex)


def add_section_band(doc: Document, text: str,
                     subtitle: str | None = None,
                     bg: str = COLOR_TITLE_BG,
                     fg: str = "FFFFFF") -> None:
    """Ajoute un bandeau de section pleine largeur (1 cellule de couleur)."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, bg)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    set_run_color(run, fg)

    if subtitle:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(2)
        r2 = p2.add_run(subtitle)
        r2.italic = True
        r2.font.size = Pt(10)
        set_run_color(r2, "DDE3EE")


def add_h1(doc: Document, text: str, color: str = COLOR_PRIMARY) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    set_run_color(run, color)


def add_h2(doc: Document, text: str, color: str = COLOR_PRIMARY_DARK) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    set_run_color(run, color)


def add_para(doc: Document, text: str, italic: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.italic = italic
    run.font.size = Pt(size)
    set_run_color(run, COLOR_TEXT)


def add_bullets(doc: Document, items: list[str], bullet_color: str = COLOR_PRIMARY) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            set_run_color(run, COLOR_TEXT)
        # Couleur du run bullet
        if not p.runs:
            run = p.add_run(item)
            run.font.size = Pt(11)
            set_run_color(run, COLOR_TEXT)
        else:
            p.runs[0].text = item
            p.runs[0].font.size = Pt(11)


def add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    """Tableau 2 colonnes (label | valeur)."""
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Light List"
    for i, (k, v) in enumerate(rows):
        c1, c2 = t.rows[i].cells
        c1.text = ""
        c2.text = ""
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(k)
        r1.bold = True
        r1.font.size = Pt(10)
        set_run_color(r1, COLOR_PRIMARY_DARK)
        p2 = c2.paragraphs[0]
        r2 = p2.add_run(v)
        r2.font.size = Pt(10)
        set_run_color(r2, COLOR_TEXT)
    # largeur colonnes
    for row in t.rows:
        row.cells[0].width = Cm(5)
        row.cells[1].width = Cm(11)


def add_etape_block(doc: Document, num: int, tag: str, titre: str,
                    description: str, objectifs: list[str],
                    competences: list[str]) -> None:
    """Bloc d'une étape : bandeau + tableau 2 colonnes."""
    # Bandeau étape
    add_section_band(
        doc,
        f"Étape {num}  ·  {tag}  ·  {titre}",
        bg=COLOR_PRIMARY,
        fg="FFFFFF",
    )

    # Tableau 2 colonnes
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_left, c_right = t.rows[0].cells
    set_cell_bg(c_left, COLOR_BLOCK_BG)
    set_cell_bg(c_right, "FFFFFF")
    c_left.width = Cm(8)
    c_right.width = Cm(8)

    # Colonne gauche : description + objectifs
    p_desc = c_left.paragraphs[0]
    rd = p_desc.add_run("Description\n")
    rd.bold = True
    rd.font.size = Pt(10)
    set_run_color(rd, COLOR_PRIMARY_DARK)

    rd2 = p_desc.add_run(description)
    rd2.font.size = Pt(10)
    set_run_color(rd2, COLOR_TEXT)

    p_obj = c_left.add_paragraph()
    rh = p_obj.add_run("Objectifs")
    rh.bold = True
    rh.font.size = Pt(10)
    set_run_color(rh, COLOR_PRIMARY_DARK)
    for obj in objectifs:
        p = c_left.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"• {obj}")
        r.font.size = Pt(10)
        set_run_color(r, COLOR_TEXT)

    # Colonne droite : compétences
    p_comp = c_right.paragraphs[0]
    rc = p_comp.add_run("Compétences acquises")
    rc.bold = True
    rc.font.size = Pt(10)
    set_run_color(rc, COLOR_SECONDARY)
    for comp in competences:
        p = c_right.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"✓ {comp}")
        r.font.size = Pt(10)
        set_run_color(r, COLOR_TEXT)


# ============================================================
# Contenu de la fiche projet
# ============================================================

PROJET = {
    "nom": "Customer Churn Analytics",
    "client_fictif": "IvoirCom",
    "secteur": "Télécom · Mobile · B2C — Côte d'Ivoire",
    "niveau": "Avancé",
    "duree": "7 à 9 heures",
    "outils": "Python, DuckDB (JupySQL), Power BI",
    "plan": "Premium (1 notebook SQL + 1 guide Power BI)",
    "ordre": "À définir",
    "publie": "Oui",
}

DESCRIPTION_GLOBALE = (
    "Customer Churn Analytics est un projet d'analyse SQL pure pour un opérateur télécom mobile "
    "ivoirien fictif (IvoirCom). L'apprenant prend le rôle d'analyste data fraîchement embauché, "
    "et doit éclairer la direction commerciale sur les départs d'abonnés. Le projet s'organise "
    "autour d'un seul notebook DuckDB / JupySQL et d'un guide Power BI complémentaire.\n\n"
    "Pas de Machine Learning : on s'appuie uniquement sur les fonctions analytiques avancées de SQL "
    "(window functions, NTILE, CTE chaînées, generate_series pour les cohortes). L'objectif est de "
    "former à la lecture business des chiffres et à la conversion d'une analyse en recommandations "
    "actionnables et chiffrées en FCFA."
)

CONTEXTE_BUSINESS = (
    "IvoirCom est un opérateur télécom mobile fictif basé à Abidjan, opérant dans 5 villes "
    "ivoiriennes (Abidjan, Bouaké, Yamoussoukro, San-Pédro, Korhogo) avec 4,2 millions d'abonnés "
    "et 6 offres au catalogue (Pulse 1 Go, Connect 5 Go, Premium 15 Go, Pro 30 Go, Étudiant 8 Go, "
    "Senior 2 Go Voix). Sa direction commerciale alerte : 12 % de la base abonnés churn chaque "
    "trimestre — soit, dans un marché où la concurrence Orange / MTN / Moov est féroce, un coût "
    "d'acquisition perdu de l'ordre de 85 000 FCFA par départ. Personne dans l'équipe n'a encore "
    "pris le temps de quantifier le phénomène par offre, par ville, par tranche d'âge, ni de "
    "vérifier si les réclamations sont un signal avant-coureur du départ. C'est ce diagnostic, "
    "et les leviers d'action chiffrés qui en découlent, que l'apprenant doit produire."
)

BRIEF_CITATION = (
    "Nos abonnés partent et on ne sait pas pourquoi. Les téléconseillers traitent les plaintes "
    "comme elles arrivent, sans priorité. Les offres Pulse et Étudiant ont des taux de désabonnement "
    "qu'on ne mesure pas. J'ai besoin d'un diagnostic clair : qui churne, pourquoi, et combien "
    "ça coûte. Et surtout, trois leviers d'action que je peux activer avec mes équipes la semaine "
    "prochaine, sans budget ML."
)
BRIEF_AUTEUR = "M. Kouadio Brou"
BRIEF_TITRE = "Directeur Commercial IvoirCom"

OBJECTIFS_PEDAGOGIQUES = [
    "Charger des CSV en tables DuckDB avec read_csv_auto et créer des vues typées (CAST en DATE, INTEGER)",
    "Identifier et corriger les anomalies de données (doublons, valeurs négatives, montants nuls, délais incohérents)",
    "Calculer des KPIs de churn (taux cumulatif, ARPU, ancienneté, écart actifs vs churners) avec COUNT FILTER, MEDIAN, DATE_DIFF",
    "Segmenter une base abonnés par offre, ville, tranche d'âge et croiser via heatmap (interaction offre × ville)",
    "Construire une analyse de cohortes avec generate_series et pivoter en heatmap de rétention 12 mois × 12 cohortes",
    "Tester l'hypothèse métier « réclamations = signal avant-coureur de churn » via window functions et UNION ALL comparatif",
    "Implémenter une segmentation RFM télécom avec NTILE(5) sur 3 dimensions (Récence facturation, Fréquence plaintes, Monétaire ARPU)",
    "Synthétiser une analyse en 3 recommandations chiffrées avec ARPU défendu en FCFA",
]

CE_QUE_TU_VAS_PRODUIRE = [
    "Un notebook SQL complet (7 sections, ~30 cellules, 18 requêtes DuckDB) avec interprétations chiffrées",
    "Une analyse de cohortes 2024 avec heatmap 12 × 12 montrant la dégradation lente vs structurelle",
    "Une segmentation RFM Telecom avec 6 segments (Champions, Loyal, At Risk, Lost, New, Others)",
    "Un dashboard Power BI 5 pages couvrant Vue Executive, Segmentation, Cohortes, Réclamations, Plan d'action",
    "Une synthèse exécutive chiffrée en 3 leviers actionnables (ARPU défendu en M FCFA sur 1 an)",
]

COMPETENCES_CV = [
    "SQL avancé : window functions (NTILE, RANK, SUM OVER PARTITION), CTE chaînées, FILTER, generate_series",
    "DuckDB : analytique en mémoire, lecture CSV native, vues typées",
    "Analyse de churn télécom : taux cumulatif, ARPU, ancienneté, signaux avant-coureurs",
    "Segmentation RFM (Récence, Fréquence, Monétaire) avec adaptation contexte télécom",
    "Power BI / DAX : KPIs métier, mesures couleur, mise en forme conditionnelle, mockup PPTX → PNG",
    "Storytelling data : conversion d'une analyse en recommandations chiffrées en FCFA",
    "Pédagogie data : cadre METHODE / INTERPRETATION / METIER appliqué à chaque analyse",
]

PREREQUIS = [
    "Python 3.10+ avec pandas, numpy, matplotlib, seaborn, duckdb, jupysql, faker (cf. requirements.txt)",
    "Power BI Desktop 2.140+ (gratuit Microsoft Store) avec extension HTML Content visual",
    "Bases SQL solides : SELECT, JOIN, GROUP BY, sous-requêtes, CTE, fenêtres analytiques",
    "Notions métier télécom : ARPU, churn, base abonnés, offres prépayées (utile mais pas indispensable)",
    "Compter 7 à 9 heures (3-4 h notebook SQL, 4-5 h guide Power BI)",
]

DONNEES_UTILISEES = [
    ("clients.csv",                 "8 030 lignes",  "Abonnés IvoirCom (id, ville, offre, date souscription, statut) — contient 30 doublons et 5 âges négatifs à corriger"),
    ("offres.csv",                  "6 lignes",      "Catalogue offres (Pulse, Connect, Premium, Pro, Étudiant, Senior) avec prix et inclus"),
    ("factures.csv",                "138 284 lignes","Facturation mensuelle 24 mois (2 % de montants nuls = 2 765 lignes à exclure de l'ARPU)"),
    ("consommation_mensuelle.csv",  "137 044 lignes","Voix / SMS / Data par client par mois (8 % de clients avec trous mensuels)"),
    ("reclamations.csv",            "9 791 lignes",  "Tickets support (3 % de délais négatifs à corriger via ABS, ~30 % de tickets en statut ouvert)"),
]

DICTIONNAIRE_DONNEES = {
    "tables": [
        {
            "name": "clients",
            "description": "Référentiel abonnés IvoirCom",
            "columns": [
                {"name": "id_client", "type": "VARCHAR", "description": "Identifiant unique abonné (CL00001 à CL08000)"},
                {"name": "nom", "type": "VARCHAR", "description": "Nom de famille (anonymisé)"},
                {"name": "prenom", "type": "VARCHAR", "description": "Prénom"},
                {"name": "sexe", "type": "VARCHAR(1)", "description": "F ou M"},
                {"name": "age", "type": "INTEGER", "description": "Âge en années (5 valeurs négatives à corriger)"},
                {"name": "ville", "type": "VARCHAR", "description": "Abidjan, Bouake, Yamoussoukro, San-Pedro, Korhogo"},
                {"name": "code_offre", "type": "VARCHAR", "description": "PULSE, CONNECT, PREMIUM, PRO, ETUDIANT, SENIOR"},
                {"name": "date_souscription", "type": "DATE", "description": "Date d'entrée chez IvoirCom"},
                {"name": "date_resiliation", "type": "DATE", "description": "NULL si actif, sinon date de churn"},
                {"name": "statut", "type": "VARCHAR", "description": "actif | resilie"},
            ],
        },
        {
            "name": "offres",
            "description": "Catalogue des 6 offres IvoirCom",
            "columns": [
                {"name": "code_offre", "type": "VARCHAR", "description": "Clé primaire offre"},
                {"name": "libelle", "type": "VARCHAR", "description": "Nom commercial (ex: Pulse 1 Go)"},
                {"name": "prix_fcfa", "type": "INTEGER", "description": "Prix mensuel en FCFA"},
                {"name": "data_inclus_go", "type": "INTEGER", "description": "Volume data inclus en Go"},
                {"name": "voix_incluses_min", "type": "INTEGER", "description": "Minutes voix incluses"},
                {"name": "segment_cible", "type": "VARCHAR", "description": "Cible marketing (Étudiants, Famille, Entrepreneurs...)"},
            ],
        },
        {
            "name": "factures",
            "description": "Facturation mensuelle 24 mois",
            "columns": [
                {"name": "id_facture", "type": "VARCHAR", "description": "Identifiant unique facture"},
                {"name": "id_client", "type": "VARCHAR", "description": "FK vers clients"},
                {"name": "date_facturation", "type": "DATE", "description": "1er du mois facturé"},
                {"name": "montant_fcfa", "type": "INTEGER", "description": "Montant total FCFA (2 % à 0 = anomalie)"},
                {"name": "methode_paiement", "type": "VARCHAR", "description": "Wave, Orange Money, MTN MoMo, Especes, Virement"},
            ],
        },
        {
            "name": "consommation_mensuelle",
            "description": "Usage voix / SMS / data par client par mois",
            "columns": [
                {"name": "id_client", "type": "VARCHAR", "description": "FK vers clients"},
                {"name": "mois", "type": "DATE", "description": "1er du mois observé"},
                {"name": "voix_min", "type": "DOUBLE", "description": "Minutes voix consommées"},
                {"name": "sms_count", "type": "INTEGER", "description": "Nombre de SMS envoyés"},
                {"name": "data_mb", "type": "DOUBLE", "description": "Mo de data consommés"},
            ],
        },
        {
            "name": "reclamations",
            "description": "Tickets support client",
            "columns": [
                {"name": "id_ticket", "type": "VARCHAR", "description": "Identifiant unique ticket"},
                {"name": "id_client", "type": "VARCHAR", "description": "FK vers clients"},
                {"name": "date_creation", "type": "DATE", "description": "Date d'ouverture du ticket"},
                {"name": "type", "type": "VARCHAR", "description": "Reseau, Facturation, Offre, Materiel"},
                {"name": "statut", "type": "VARCHAR", "description": "resolu, ouvert, abandonne"},
                {"name": "delai_resolution_jours", "type": "INTEGER", "description": "NULL si ouvert, négatif = anomalie à corriger"},
            ],
        },
    ]
}

REFLEXION_STRATEGIQUE = [
    "Comment formaliser le « churn » dans une base abonnés télécom — par date de résiliation explicite ou par absence de facturation sur N mois ?",
    "Pourquoi distinguer le taux de churn cumulatif (sur toute la période) du taux trimestriel ? Lequel est le plus parlant pour un Comité Exécutif ?",
    "L'ARPU moyen est-il toujours pertinent quand la distribution est asymétrique (queue longue côté offres Pro) ? Quand préférer la médiane ?",
    "Une analyse de cohortes peut-elle révéler un événement opérationnel (panne réseau, lancement concurrent) — comment ?",
    "Quel est le seuil critique au-delà duquel le segment At Risk doit déclencher une cellule rétention dédiée ? Comment le justifier en FCFA ?",
    "Les fonctions NTILE produisent des quintiles équilibrés en taille — est-ce toujours souhaitable, ou faut-il parfois des seuils métier fixes (R = 30j / 60j / 90j) ?",
    "Comment combiner un score RFM avec un signal opérationnel (« 2+ tickets non résolus ») pour prioriser les appels sortants de demain matin ?",
]

# ============================================================
# Définition des étapes (= notebooks)
# ============================================================

ETAPES = [
    {
        "num": 1,
        "nom": "sql-analytics-churn",
        "nom_court": "SQL Analytics",
        "tag": "NOTEBOOK 1",
        "titre": "SQL Analytics, KPIs & Segmentation Churn",
        "description": (
            "Tu prends le rôle d'analyste data chez IvoirCom. À partir des 5 CSV bruts (clients, offres, "
            "factures, consommation, réclamations), tu construis un diagnostic churn complet en SQL pur : "
            "nettoyage des anomalies, KPIs globaux, segmentation par offre/ville/âge, analyse de cohortes 2024, "
            "test des signaux faibles via les réclamations, segmentation RFM Telecom à 6 segments. "
            "Tu termines en rédigeant 3 recommandations chiffrées en FCFA pour la direction commerciale."
        ),
        "objectifs": [
            "Charger 5 CSV en tables DuckDB et créer des vues typées (CAST en DATE, INTEGER)",
            "Détecter et corriger 4 types d'anomalies : doublons _DUP, âges négatifs, montants nuls, délais négatifs",
            "Calculer le taux de churn cumulatif global et par segment (offre, ville, âge)",
            "Construire une heatmap de rétention 12 mois × 12 cohortes via generate_series",
            "Tester le signal '2+ tickets non résolus' et chiffrer son écart vs base globale",
            "Implémenter un RFM Telecom avec NTILE(5) et 6 segments métier",
            "Rédiger une synthèse exécutive avec 3 leviers chiffrés en FCFA",
        ],
        "competences": [
            "DuckDB + JupySQL : connexion, magic %%sql, vues typées",
            "Window functions : NTILE, SUM OVER PARTITION, RANKX",
            "CTE chaînées multi-niveaux",
            "generate_series pour expansion temporelle",
            "Heatmap matplotlib / seaborn (rétention, croisement offre × ville)",
            "Segmentation RFM adaptée au contexte télécom",
            "Storytelling data : METHODE / INTERPRETATION / METIER",
        ],
    },
    {
        "num": 2,
        "nom": "dashboard-powerbi",
        "nom_court": "Power BI",
        "tag": "GUIDE POWER BI",
        "titre": "Dashboard de pilotage Customer Churn — 5 pages",
        "description": (
            "À partir des CSV nettoyés (mêmes sources que le notebook SQL), tu construis dans Power BI Desktop "
            "un dashboard de pilotage pour la direction commerciale d'IvoirCom. Le dashboard couvre 5 pages : "
            "Vue Executive (KPIs globaux + tendance churn), Segments à risque (offres / villes / âges), "
            "Cohortes & Rétention (heatmap), Signaux Réclamations (funnel statut + ratio churn par type plainte), "
            "Plan d'action (RFM segmenté + Top 50 At Risk). Le guide pas-à-pas couvre la modélisation en étoile, "
            "les ~40 mesures DAX, la charte Orange chaleureux, et le mockup PPTX → PNG."
        ),
        "objectifs": [
            "Charger les 5 CSV via Power Query avec lecture GitHub raw (reproductibilité)",
            "Créer la table Calendrier et la marquer comme table de dates",
            "Construire un modèle en étoile avec relations actives sur clients/factures/réclamations",
            "Implémenter ~40 mesures DAX (taux churn %, ARPU, RFM, color measures, alertes)",
            "Construire 5 pages avec navbar horizontale, slicers synchronisés, charte Orange chaleureux",
            "Produire un mockup PPTX vierge → 5 PNG arrière-plans à 150 DPI",
            "Valider la recette finale (checklist : 5 relations actives, 0 LocalDateTable, formats DAX corrects)",
        ],
        "competences": [
            "Power BI Desktop : Power Query, modélisation, mesures DAX, mise en page",
            "DAX avancé : VAR/RETURN, RANKX, DATEADD, DATESYTD, LOOKUPVALUE, CALCULATE + FILTER",
            "Mesures couleur dynamiques pour mise en forme conditionnelle",
            "HTML Content visual marketplace (Daniel Marsh-Patrick) pour heatmaps custom",
            "Mockup PowerPoint vierge → export PNG 150 DPI → arrière-plan Power BI",
            "Checklist de recette dashboard professionnel",
            "Storytelling exécutif en 5 pages avec navigation",
        ],
    },
]

RESSOURCES = [
    ("dataset/clients.csv",                       "CSV", "Référentiel abonnés (8 030 lignes brutes)"),
    ("dataset/offres.csv",                        "CSV", "Catalogue 6 offres"),
    ("dataset/factures.csv",                      "CSV", "Facturation mensuelle (138 284 lignes)"),
    ("dataset/consommation_mensuelle.csv",        "CSV", "Usage voix/SMS/data (137 044 lignes)"),
    ("dataset/reclamations.csv",                  "CSV", "Tickets support (9 791 lignes)"),
    ("notebooks/enonce/Notebook_1_SQL_Analytics_Churn.ipynb",
                                                  "Notebook", "Énoncé apprenant — cellules vides + TODO + questions"),
    ("notebooks/solution/Notebook_1_SQL_Analytics_Churn_--_Solution.ipynb",
                                                  "Notebook", "Corrigé complet avec interprétations chiffrées"),
    ("powerbi/ressources_powerbi.ipynb",          "Notebook", "Guide pas-à-pas Power BI (modélisation, DAX, 5 pages)"),
    ("powerbi/mockup_customer_churn.pptx",        "PPTX",     "Mockup vierge 5 slides (1280×720, charte Orange)"),
    ("powerbi/bg-XX-*.png",                       "PNG",      "5 PNG arrière-plans Power BI (150 DPI)"),
    ("insights/recommandations_churn.docx",       "DOCX",     "Synthèse exécutive 3 recommandations chiffrées"),
    ("generate_datasets.py",                      "Script",   "Génération reproductible des 5 CSV (Faker + logique métier churn)"),
    ("requirements.txt",                          "Texte",    "Dépendances Python du projet"),
]

# ============================================================
# Construction du document
# ============================================================

def build_doc() -> Document:
    doc = Document()

    # Marges
    for section in doc.sections:
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

    # Style par défaut
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ----- BANDEAU TITRE -----
    add_section_band(
        doc,
        f"DataProjectLab  ·  {PROJET['nom']}",
        subtitle="Fiche projet — Création sur plateforme",
        bg=COLOR_TITLE_BG,
        fg="FFFFFF",
    )

    doc.add_paragraph()

    # ----- TABLEAU FICHE PROJET -----
    add_h1(doc, "Fiche projet")
    add_kv_table(doc, [
        ("Nom du projet",      PROJET["nom"]),
        ("Client fictif",      PROJET["client_fictif"]),
        ("Secteur d'activité", PROJET["secteur"]),
        ("Niveau",             PROJET["niveau"]),
        ("Durée estimée",      PROJET["duree"]),
        ("Outils utilisés",    PROJET["outils"]),
        ("Plan minimum",       PROJET["plan"]),
        ("Ordre d'affichage",  PROJET["ordre"]),
        ("Publié",             PROJET["publie"]),
    ])

    doc.add_paragraph()

    # ----- DESCRIPTION GLOBALE -----
    add_h1(doc, "Description du projet")
    for para in DESCRIPTION_GLOBALE.split("\n\n"):
        add_para(doc, para)

    # ----- MISSION / CONTEXTE BUSINESS -----
    add_h1(doc, "Mission — Contexte business")
    add_para(doc, CONTEXTE_BUSINESS)

    # ----- BRIEF METIER -----
    add_h1(doc, "Brief métier — Citation")

    # Encadre principal : citation + auteur + titre dans un tableau structure
    brief_table = doc.add_table(rows=3, cols=2)
    brief_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    brief_table.autofit = False

    # Ligne 1 : Citation (cellules fusionnees sur 2 colonnes)
    citation_cell = brief_table.rows[0].cells[0]
    citation_cell.merge(brief_table.rows[0].cells[1])
    set_cell_bg(citation_cell, COLOR_BLOCK_BG)
    p_label_cit = citation_cell.paragraphs[0]
    r_label_cit = p_label_cit.add_run("Texte du brief métier")
    r_label_cit.bold = True
    r_label_cit.font.size = Pt(9)
    set_run_color(r_label_cit, COLOR_PRIMARY_DARK)
    p_cit = citation_cell.add_paragraph()
    p_cit.paragraph_format.left_indent = Cm(0.2)
    r_cit = p_cit.add_run(f"« {BRIEF_CITATION} »")
    r_cit.italic = True
    r_cit.font.size = Pt(11)
    set_run_color(r_cit, COLOR_TEXT)

    # Ligne 2 : Labels Auteur / Titre
    label_auteur_cell = brief_table.rows[1].cells[0]
    label_titre_cell = brief_table.rows[1].cells[1]
    set_cell_bg(label_auteur_cell, "FFFFFF")
    set_cell_bg(label_titre_cell, "FFFFFF")
    for cell, label in [(label_auteur_cell, "Auteur"),
                        (label_titre_cell, "Titre")]:
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(9)
        set_run_color(r, COLOR_PRIMARY_DARK)

    # Ligne 3 : Valeurs Auteur / Titre
    val_auteur_cell = brief_table.rows[2].cells[0]
    val_titre_cell = brief_table.rows[2].cells[1]
    set_cell_bg(val_auteur_cell, "FFFFFF")
    set_cell_bg(val_titre_cell, "FFFFFF")
    for cell, value in [(val_auteur_cell, BRIEF_AUTEUR),
                        (val_titre_cell, BRIEF_TITRE)]:
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(value)
        r.font.size = Pt(11)
        set_run_color(r, COLOR_TEXT)

    # Largeurs egales 8 cm chacune
    for row in brief_table.rows:
        for c in row.cells:
            c.width = Cm(8)

    # ----- OBJECTIFS PEDAGOGIQUES -----
    add_h1(doc, "Objectifs pédagogiques")
    add_bullets(doc, OBJECTIFS_PEDAGOGIQUES)

    # ----- CE QUE TU VAS PRODUIRE -----
    add_h1(doc, "Ce que tu vas produire")
    add_bullets(doc, CE_QUE_TU_VAS_PRODUIRE)

    # ----- COMPETENCES CV -----
    add_h1(doc, "Compétences à mettre sur ton CV")
    add_bullets(doc, COMPETENCES_CV)

    # ----- AVANT DE DEMARRER -----
    add_h1(doc, "Avant de démarrer — Prérequis")
    add_bullets(doc, PREREQUIS)

    # ----- DONNEES UTILISEES -----
    add_h1(doc, "Données utilisées")
    t = doc.add_table(rows=1 + len(DONNEES_UTILISEES), cols=3)
    t.style = "Light Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(["Fichier", "Volumétrie", "Description"]):
        c = hdr[i]
        set_cell_bg(c, COLOR_PRIMARY)
        c.text = ""
        rh = c.paragraphs[0].add_run(h)
        rh.bold = True
        rh.font.size = Pt(10)
        set_run_color(rh, "FFFFFF")
    for i, (f, n, d) in enumerate(DONNEES_UTILISEES, start=1):
        cells = t.rows[i].cells
        for j, val in enumerate([f, n, d]):
            cells[j].text = ""
            r = cells[j].paragraphs[0].add_run(val)
            r.font.size = Pt(9)
            set_run_color(r, COLOR_TEXT)
        cells[0].width = Cm(5)
        cells[1].width = Cm(2.5)
        cells[2].width = Cm(8.5)

    # ----- DICTIONNAIRE DES DONNEES (JSON) -----
    add_h1(doc, "Dictionnaire des données (JSON)")
    add_para(doc,
             "Le bloc ci-dessous contient le dictionnaire complet des 5 tables au format JSON. "
             "Il peut être copié-collé tel quel dans le champ « Dictionnaire des données » de la "
             "page projet sur la plateforme DataProjectLab.",
             italic=True, size=10)

    p_json = doc.add_paragraph()
    p_json.paragraph_format.left_indent = Cm(0.3)
    r_json = p_json.add_run(json.dumps(DICTIONNAIRE_DONNEES, indent=2, ensure_ascii=False))
    r_json.font.name = "Consolas"
    r_json.font.size = Pt(8)
    set_run_color(r_json, COLOR_TEXT)

    # ----- REFLEXION STRATEGIQUE -----
    add_h1(doc, "Réflexion stratégique")
    add_bullets(doc, REFLEXION_STRATEGIQUE)

    # ----- ETAPES DU PROJET -----
    doc.add_page_break()
    add_section_band(
        doc,
        f"Étapes du projet",
        subtitle=f"{len(ETAPES)} étapes structurées — Du diagnostic SQL au dashboard de pilotage",
        bg=COLOR_TITLE_BG,
        fg="FFFFFF",
    )

    for etape in ETAPES:
        doc.add_paragraph()
        add_etape_block(
            doc,
            num=etape["num"],
            tag=etape["tag"],
            titre=etape["titre"],
            description=etape["description"],
            objectifs=etape["objectifs"],
            competences=etape["competences"],
        )

        # Métadonnées plateforme
        doc.add_paragraph()
        add_h2(doc, f"Métadonnées plateforme — Étape {etape['num']}")
        add_kv_table(doc, [
            ("Nom (slug kebab-case)", etape["nom"]),
            ("Nom court",             etape["nom_court"]),
            ("Tag",                   etape["tag"]),
            ("Titre",                 etape["titre"]),
        ])

    # ----- RESSOURCES -----
    doc.add_paragraph()
    add_h1(doc, "Ressources livrées avec le projet")
    t = doc.add_table(rows=1 + len(RESSOURCES), cols=3)
    t.style = "Light Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(["Chemin / Nom", "Type", "Description"]):
        c = hdr[i]
        set_cell_bg(c, COLOR_PRIMARY)
        c.text = ""
        rh = c.paragraphs[0].add_run(h)
        rh.bold = True
        rh.font.size = Pt(10)
        set_run_color(rh, "FFFFFF")
    for i, (path, typ, desc) in enumerate(RESSOURCES, start=1):
        cells = t.rows[i].cells
        for j, val in enumerate([path, typ, desc]):
            cells[j].text = ""
            r = cells[j].paragraphs[0].add_run(val)
            r.font.size = Pt(9)
            set_run_color(r, COLOR_TEXT)
        cells[0].width = Cm(7)
        cells[1].width = Cm(2)
        cells[2].width = Cm(7)

    # ----- FOOTER -----
    doc.add_paragraph()
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_foot = p_foot.add_run(
        "DataProjectLab — apprendre la data sur des cas concrets, "
        "structurés et orientés métier."
    )
    r_foot.italic = True
    r_foot.font.size = Pt(9)
    set_run_color(r_foot, COLOR_LABEL)

    return doc


def main() -> None:
    out = Path(__file__).parent / "fiche_projet_customer_churn_analytics.docx"
    doc = build_doc()
    doc.save(out)
    print(f"[OK] Fiche projet generee : {out}")
    print(f"     Taille : {out.stat().st_size / 1024:.1f} Ko")


if __name__ == "__main__":
    main()
