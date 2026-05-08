"""
build_recommandations_edutrack.py
Genere le rapport de synthese EduTrack en .docx (5 pages max).

Sortie : recommandations_edutrack.docx (meme dossier que ce script)

Prerequis :
    pip install python-docx requests

Usage :
    python build_recommandations_edutrack.py
"""

from docx import Document
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from pathlib import Path
import urllib.request
import sys

ROOT = Path(__file__).parent
OUT = ROOT / "recommandations_edutrack.docx"
LOGO_URL = "https://raw.githubusercontent.com/dataprojectlabs/DataProjectLab-projects/refs/heads/main/media/logo_dataprojectlab.png"
LOGO_CACHE = ROOT / ".cache_logo.png"

# Charte DataProjectLab
NAVY = RGBColor(0x1E, 0x3A, 0x5F)
VIOLET = RGBColor(0x53, 0x4A, 0xB7)
VERT = RGBColor(0x1F, 0xA6, 0x7D)
ROUGE = RGBColor(0xE5, 0x49, 0x4D)
ORANGE = RGBColor(0xF2, 0xA9, 0x3B)
GRIS = RGBColor(0x6B, 0x72, 0x80)
TEXTE = RGBColor(0x11, 0x18, 0x27)
FOND_PALE = "EEEDFE"
FOND_GREEN = "DCFCE7"
FOND_RED = "FEE2E2"
FOND_ORANGE = "FEF3C7"
FOND_NAVY = "1E3A5F"


# ============================================================
# Helpers
# ============================================================

def download_logo():
    if not LOGO_CACHE.exists():
        urllib.request.urlretrieve(LOGO_URL, LOGO_CACHE)
    return LOGO_CACHE


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, color="CCCCCC", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:color'), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_para(doc, text, *, size=11, bold=False, color=TEXTE, align=None,
             space_before=0, space_after=4, italic=False, font="Calibri"):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    return p


def add_heading(doc, text, *, size=16, color=NAVY, space_before=12, space_after=6, font="Georgia"):
    return add_para(doc, text, size=size, bold=True, color=color,
                    space_before=space_before, space_after=space_after, font=font)


def add_bullet(doc, text, *, size=10, color=TEXTE):
    p = doc.add_paragraph(style="List Bullet")
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def add_kpi_row(doc, kpis):
    """kpis = [(value, label, hex_color), ...]"""
    table = doc.add_table(rows=2, cols=len(kpis))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (val, lbl, color) in enumerate(kpis):
        v_cell = table.rows[0].cells[i]
        l_cell = table.rows[1].cells[i]
        v_cell.text = ""
        l_cell.text = ""
        v_p = v_cell.paragraphs[0]
        v_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        v_run = v_p.add_run(val)
        v_run.font.name = "Georgia"
        v_run.font.size = Pt(20)
        v_run.font.bold = True
        v_run.font.color.rgb = color
        l_p = l_cell.paragraphs[0]
        l_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        l_run = l_p.add_run(lbl)
        l_run.font.name = "Calibri"
        l_run.font.size = Pt(9)
        l_run.font.color.rgb = GRIS
    # Pas de bordures
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = OxmlElement('w:tcBorders')
            for edge in ('top', 'left', 'bottom', 'right'):
                b = OxmlElement(f'w:{edge}')
                b.set(qn('w:val'), 'nil')
                tc_borders.append(b)
            tc_pr.append(tc_borders)


def add_constat_card(doc, titre, contenu, accent_hex="534AB7"):
    """Encadre type card avec border-top accent."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.text = ""
    set_cell_shading(cell, "F9F9F8")
    # Border-top accent
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '24')
    top.set(qn('w:color'), accent_hex)
    tc_borders.append(top)
    for edge in ('left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), 'E5E7EB')
        tc_borders.append(b)
    tc_pr.append(tc_borders)
    # Contenu
    p_titre = cell.paragraphs[0]
    pf = p_titre.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(2)
    run = p_titre.add_run(titre)
    run.font.name = "Georgia"
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = NAVY
    # Corps
    p_corps = cell.add_paragraph()
    pf = p_corps.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(4)
    run = p_corps.add_run(contenu)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = TEXTE


# ============================================================
# Construction
# ============================================================

doc = Document()

# Marges (A4)
for section in doc.sections:
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# ============================================================
# PAGE 1 — Couverture
# ============================================================

# Bandeau navy avec logo
logo_path = download_logo()
header_table = doc.add_table(rows=1, cols=2)
header_table.autofit = False
hcell_l = header_table.rows[0].cells[0]
hcell_r = header_table.rows[0].cells[1]
hcell_l.width = Cm(11)
hcell_r.width = Cm(5)
set_cell_shading(hcell_l, FOND_NAVY)
set_cell_shading(hcell_r, FOND_NAVY)

# Texte gauche
hcell_l.text = ""
p_tag = hcell_l.paragraphs[0]
pf = p_tag.paragraph_format
pf.space_before = Pt(28); pf.space_after = Pt(4)
r = p_tag.add_run("DATAPROJECTLAB · POWER BI · EDTECH")
r.font.name = "Calibri"; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x9A, 0xB6, 0xE0)
p_titre = hcell_l.add_paragraph()
pf = p_titre.paragraph_format; pf.space_before = Pt(2); pf.space_after = Pt(0)
r = p_titre.add_run("EduTrack Analytics")
r.font.name = "Georgia"; r.font.size = Pt(28); r.font.bold = True
r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
p_sub = hcell_l.add_paragraph()
pf = p_sub.paragraph_format; pf.space_before = Pt(2); pf.space_after = Pt(28)
r = p_sub.add_run("Rapport de synthèse — Détection du décrochage")
r.font.name = "Calibri"; r.font.size = Pt(12); r.font.italic = True
r.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE0)

# Logo droite
hcell_r.text = ""
p_logo = hcell_r.paragraphs[0]
p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf = p_logo.paragraph_format; pf.space_before = Pt(38)
r = p_logo.add_run()
r.add_picture(str(logo_path), width=Cm(3.2))

# Pas de bordures
for row in header_table.rows:
    for cell in row.cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right'):
            b = OxmlElement(f'w:{edge}'); b.set(qn('w:val'), 'nil')
            tc_borders.append(b)
        tc_pr.append(tc_borders)

# Espace
add_para(doc, "", size=10, space_after=12)

# Métadonnées
add_para(doc, "Document destiné à la direction pédagogique d'EduTrack",
         size=11, italic=True, color=GRIS, space_after=4)
add_para(doc, "Date : Mai 2026   ·   Version 1.0   ·   Préparé par DataProjectLab Academy",
         size=10, color=GRIS, space_after=20)

# Sous-titre
add_heading(doc, "Chiffres clés", size=14, space_before=8, space_after=10)

# 4 KPIs
add_kpi_row(doc, [
    ("4 500", "Apprenants actifs", VIOLET),
    ("33,3 %", "Taux de complétion", VERT),
    ("17,7 %", "Taux d'abandon", ROUGE),
    ("778 M FCFA", "Revenu généré", VIOLET),
])
add_para(doc, "", size=8, space_after=8)
add_kpi_row(doc, [
    ("1 613", "Apprenants en alerte ML", ROUGE),
    ("4,23 / 5", "CSAT moyen", ORANGE),
    ("12 parcours", "Catalogue actif", VIOLET),
    ("79,6 %", "Recall modèle ML", VERT),
])

# Page break
doc.add_page_break()

# ============================================================
# PAGE 2 — Synthèse exécutive + Constats
# ============================================================

add_heading(doc, "1.  Synthèse exécutive", size=15, space_before=0, space_after=8)

add_para(doc,
    "EduTrack Academy a accueilli 4 500 apprenants et généré 778 M FCFA "
    "(1,19 M EUR) sur les 12 derniers mois, avec une progression de +21 % "
    "des inscriptions vs N-1. Mais le taux de complétion reste à 33,3 %, "
    "soit 7 points sous la cible interne de 40 %.",
    size=11, space_after=6)

add_para(doc,
    "Trois enjeux ressortent du modèle de risque ML : 1 613 apprenants "
    "(36 % du parc actif) sont en alerte de décrochage, dont 150 en zone "
    "critique nécessitant une intervention humaine immédiate. La qualité "
    "pédagogique varie fortement : 4 instructeurs sur 12 sont classés "
    "« À coacher » (faible complétion ET CSAT en dessous de la médiane).",
    size=11, space_after=6)

add_para(doc,
    "Trois leviers sont disponibles à 6 mois : (1) coaching ciblé des "
    "instructeurs sous-performants, (2) séquence d'engagement automatisée "
    "sur les apprenants à risque, (3) capitalisation sur les Top 5 parcours "
    "qui portent 60 % du revenu.",
    size=11, space_after=14)

add_heading(doc, "2.  Constats data-driven", size=15, space_before=4, space_after=8)

add_constat_card(doc,
    "Pédagogie — un écart structurel à la cible",
    "Taux de complétion 33,3 % vs cible 40 % (-7 pp). Taux d'abandon 17,7 %, "
    "stable depuis 12 mois. CSAT 4,23 / 5 (correct mais polarisé : 4 instructeurs "
    "tirent la moyenne vers le bas). Pire parcours par domaine : Machine Learning "
    "Appliqué (Data & IA, 16,9 % d'abandon).",
    accent_hex="534AB7")

add_para(doc, "", size=8, space_after=4)

add_constat_card(doc,
    "Risque ML — 1 613 apprenants à mobiliser",
    "Le modèle a identifié 1 613 apprenants à risque sur 4 500 actifs. "
    "Répartition : ~150 ROUGE (score > 0,80, action urgente), ~150 ORANGE "
    "(0,60-0,80, relance ciblée), ~1 300 VIOLET (0,30-0,60, rappel doux). "
    "Recall 79,6 % : sur 100 apprenants qui décrocheront réellement, le modèle "
    "en identifie 80.",
    accent_hex="E5494D")

add_para(doc, "", size=8, space_after=4)

add_constat_card(doc,
    "Business — concentration sur le Top 5",
    "Les 5 premiers parcours (React & Node.js, Machine Learning, Python Django, "
    "Data Python, Power BI) génèrent 60 % du revenu. Mobile Money domine les "
    "paiements (50 %), suivi de la Carte (26 %). Côte d'Ivoire = 40 % du parc, "
    "les 23-32 ans représentent 60 %.",
    accent_hex="1FA67D")

doc.add_page_break()

# ============================================================
# PAGE 3 — 3 Recommandations chiffrées
# ============================================================

add_heading(doc, "3.  Recommandations prioritaires", size=15, space_before=0, space_after=10)

# Tableau 3 recommandations × 4 colonnes
reco_data = [
    ("R1 — Plan de coaching instructeurs",
     "Identifier les 4 instructeurs en quadrant « À coacher » et « À accompagner ». "
     "Programme de 6 semaines : revue pédagogique mensuelle, observation de classe, "
     "refonte du syllabus du parcours qui décroche le plus.",
     "+2 pp de complétion (33,3 % → 35,3 %)\n6 semaines\nCoût estimé : 8 M FCFA",
     "534AB7"),
    ("R2 — Séquence d'engagement ML",
     "Activer 3 parcours d'action selon le niveau d'alerte : appel téléphonique "
     "pour les ~150 ROUGE, email personnalisé pour les ~150 ORANGE, push de "
     "réengagement pour les ~1 300 VIOLET. Mesurer le re-engagement à 7 et 30 jours.",
     "-3 pp d'abandon (17,7 % → 14,7 %)\n3 mois\nCoût estimé : 4 M FCFA",
     "E5494D"),
    ("R3 — Capitalisation Top 5 + refonte des pires parcours",
     "Doubler la capacité d'inscription sur les Top 5 (campagne acquisition + "
     "places supplémentaires). Auditer puis refondre les 4 pires parcours par "
     "domaine (Machine Learning Appliqué, React & Node.js Fullstack, SEO & Growth, "
     "Management de Projet).",
     "+15 % de revenu (778 M → 895 M FCFA)\n6 mois\nROI : x4",
     "1FA67D"),
]

table = doc.add_table(rows=len(reco_data) + 1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
col_widths = [Cm(4.2), Cm(7.3), Cm(3.8), Cm(0.8)]  # le 0.8 sert juste de marge
for i, w in enumerate(col_widths[:3]):
    table.columns[i].width = w

# Header
header_cells = table.rows[0].cells
for h_cell in header_cells[:3]:
    set_cell_shading(h_cell, FOND_NAVY)
    set_cell_borders(h_cell, "1E3A5F", "4")
header_cells[0].text = ""
header_cells[1].text = ""
header_cells[2].text = ""
for i, lbl in enumerate(["Recommandation", "Action", "Impact attendu"]):
    p = header_cells[i].paragraphs[0]
    r = p.add_run(lbl)
    r.font.name = "Calibri"; r.font.size = Pt(10); r.font.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)

# Lignes
for i, (titre, action, impact, accent) in enumerate(reco_data):
    row = table.rows[i + 1]
    c0, c1, c2 = row.cells[0], row.cells[1], row.cells[2]
    for c in (c0, c1, c2):
        set_cell_borders(c, "E5E7EB", "4")
        # Border left accent
        tc_pr = c._tc.get_or_add_tcPr()
        for child in list(tc_pr):
            if child.tag == qn('w:tcBorders'):
                tc_pr.remove(child)
        tc_borders = OxmlElement('w:tcBorders')
        for edge in ('top', 'bottom', 'right'):
            b = OxmlElement(f'w:{edge}')
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), '4')
            b.set(qn('w:color'), 'E5E7EB')
            tc_borders.append(b)
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        if c is c0:
            left.set(qn('w:sz'), '24')
            left.set(qn('w:color'), accent)
        else:
            left.set(qn('w:sz'), '4')
            left.set(qn('w:color'), 'E5E7EB')
        tc_borders.append(left)
        tc_pr.append(tc_borders)

    # Titre R
    c0.text = ""
    p = c0.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(titre)
    r.font.name = "Georgia"; r.font.size = Pt(11); r.font.bold = True
    r.font.color.rgb = NAVY

    # Action
    c1.text = ""
    p = c1.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(action)
    r.font.name = "Calibri"; r.font.size = Pt(10); r.font.color.rgb = TEXTE

    # Impact
    c2.text = ""
    p = c2.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    for j, line in enumerate(impact.split("\n")):
        if j > 0:
            p = c2.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
        r = p.add_run(line)
        r.font.name = "Calibri"; r.font.size = Pt(9.5)
        r.font.bold = (j == 0)
        r.font.color.rgb = NAVY if j == 0 else GRIS

doc.add_page_break()

# ============================================================
# PAGE 4 — Limites & Prochaines étapes
# ============================================================

add_heading(doc, "4.  Limites du modèle et des données", size=15, space_before=0, space_after=8)

for txt in [
    "Recall modèle ML : 79,6 %. Sur 100 apprenants qui décrochent réellement, le modèle en flag 80 — les 20 manqués constituent une zone aveugle.",
    "Le score de risque s'appuie sur des features comportementales (jours d'inactivité, engagement, progression) mais n'intègre pas le contexte personnel (charge professionnelle, événements de vie).",
    "Saisonnalité partiellement intégrée : la baisse de juin n'est expliquée qu'à 60 % par les variables disponibles.",
    "Pas de feedback utilisateur post-abandon : on ne sait pas pourquoi un apprenant arrête. Une enquête qualitative à 30 jours après un abandon permettrait d'affiner les actions de rétention.",
    "Modèle entraîné sur 24 mois d'historique : à réentraîner trimestriellement pour suivre l'évolution des comportements.",
]:
    add_bullet(doc, txt, size=10)

add_para(doc, "", size=8, space_after=8)

add_heading(doc, "5.  Prochaines étapes", size=15, space_before=4, space_after=8)

roadmap = [
    ("À 1 mois", "1FA67D",
     "Lancer la séquence ROUGE (150 appels téléphoniques). Cadrer le programme de coaching avec les 4 instructeurs. Mesurer baseline avant intervention."),
    ("À 3 mois", "F2A93B",
     "Activer ORANGE et VIOLET. Premier point d'étape coaching (revue pédagogique). Évaluer l'effet sur le taux d'abandon mensuel."),
    ("À 6 mois", "534AB7",
     "Bilan complet des 3 recommandations : impact sur complétion (cible +2 pp), abandon (cible -3 pp), revenu (cible +15 %). Décision de poursuite ou ajustement."),
]
for label, color, txt in roadmap:
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(2.8)
    table.columns[1].width = Cm(13.2)
    cell_l = table.rows[0].cells[0]
    cell_r = table.rows[0].cells[1]
    set_cell_shading(cell_l, color)
    cell_l.text = ""
    p = cell_l.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(label)
    r.font.name = "Georgia"; r.font.size = Pt(11); r.font.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    cell_r.text = ""
    p = cell_r.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.4)
    r = p.add_run(txt)
    r.font.name = "Calibri"; r.font.size = Pt(10); r.font.color.rgb = TEXTE
    # Pas de bordures
    for c in (cell_l, cell_r):
        tc_pr = c._tc.get_or_add_tcPr()
        tc_borders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right'):
            b = OxmlElement(f'w:{edge}'); b.set(qn('w:val'), 'nil')
            tc_borders.append(b)
        tc_pr.append(tc_borders)
    add_para(doc, "", size=4, space_after=2)

doc.add_page_break()

# ============================================================
# PAGE 5 — Annexe
# ============================================================

add_heading(doc, "Annexe — Glossaire des métriques clés", size=14, space_before=0, space_after=8)

glossaire = [
    ("Taux de complétion", "Pourcentage d'inscriptions ayant atteint le statut « Termine » sur le total des inscriptions. Cible interne EduTrack : 40 %."),
    ("Taux d'abandon", "Pourcentage d'inscriptions ayant atteint le statut « Abandonne ». Stable à 17,7 % sur 12 mois."),
    ("CSAT", "Customer Satisfaction Score — note moyenne (1 à 5) attribuée par les apprenants en fin de parcours."),
    ("Quadrant Instructeur", "Classement croisant Taux de complétion (médiane 33,3 %) et CSAT (médiane 4,23) en 4 segments : Top Performer, À accompagner, Coach CSAT, À coacher."),
    ("Score de risque ML", "Probabilité (0-1) qu'un apprenant abandonne, prédite par un modèle de classification entraîné sur 24 mois d'historique."),
    ("Niveau d'alerte", "Discrétisation du score : ROUGE (>0,80, intervention humaine), ORANGE (0,60-0,80, relance ciblée), VIOLET (0,30-0,60, rappel doux), VERT (<0,30, RAS)."),
    ("Recall", "Sur 100 apprenants qui abandonneront réellement, combien le modèle en a-t-il identifiés ? EduTrack : 79,6 %."),
    ("RevPAR", "Revenue per Available Spot — revenu moyen par place ouverte sur la plateforme. Indicateur de rentabilité par parcours."),
]
for terme, defn in glossaire:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(2); pf.space_after = Pt(4)
    r = p.add_run(terme + " — ")
    r.font.name = "Calibri"; r.font.size = Pt(10); r.font.bold = True
    r.font.color.rgb = NAVY
    r2 = p.add_run(defn)
    r2.font.name = "Calibri"; r2.font.size = Pt(10); r2.font.color.rgb = TEXTE

# Footer tagline
add_para(doc, "", size=10, space_after=12)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("DataProjectLab — apprendre la data sur des cas concrets, structurés et orientés métier.")
r.font.name = "Calibri"; r.font.size = Pt(9); r.font.italic = True
r.font.color.rgb = GRIS


# ============================================================
# Sauvegarde
# ============================================================

doc.save(str(OUT))
print(f"\n[OK] Rapport genere : {OUT}")
print(f"     Taille : {OUT.stat().st_size / 1024:.1f} ko")
